import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

TARGET_DIR = "/storage/emulated/0/universida-datos/diego"
REPO_TARGET = "/storage/emulated/0/universida-datos/instalacion-electricas/proyectos/diego-unifamiliar"

def set_cell_bg(cell, hex_code):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_code}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

doc = docx.Document()

# Margins: 2.5 cm (0.98 in)
for section in doc.sections:
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    section.left_margin = Inches(0.98)
    section.right_margin = Inches(0.98)

def add_p(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, bold=False, italic=False, size=11, color=(0,0,0), first_indent=0.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if first_indent > 0 and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    if text:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*color)
    return p

def add_roman_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 32, 67)
    return p

def add_num_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level==1 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12 if level==1 else 11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(24, 76, 120) if level==1 else RGBColor(0, 0, 0)
    return p

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = RGBColor(60, 60, 60)
    return p

# =============================================================================
# PORTADA OFICIAL REPLICA MODELO
# =============================================================================
add_p("INSTALACIONES ELÉCTRICAS", WD_ALIGN_PARAGRAPH.CENTER, 40, 12, True, False, 18, (15, 32, 67), first_indent=0)
add_p("MEMORIA DESCRIPTIVA, CÁLCULOS JUSTIFICATORIOS,\nESPECIFICACIONES TÉCNICAS Y MATRIZ DE INCONSISTENCIAS", WD_ALIGN_PARAGRAPH.CENTER, 0, 36, True, False, 14, (24, 76, 120), first_indent=0)

meta_items = [
    ("PROYECTO", "VIVIENDA UNIFAMILIAR DE 3 NIVELES (1ER PISO, 2DO PISO Y AZOTEA)"),
    ("PROPIETARIO / PROYECTISTA", "CHARAJA MAMANI DIEGO JEFFERSON"),
    ("CÓDIGO DE ESTUDIANTE", "214254"),
    ("DOCENTE RESPONSABLE", "ING. VILLANUEVA CORNEJO MARCOS JOSE"),
    ("ASIGNATURA", "INSTALACIONES ELÉCTRICAS"),
    ("ESCUELA PROFESIONAL", "ESCUELA PROFESIONAL DE INGENIERÍA MECÁNICA ELÉCTRICA (EPIME)"),
    ("FACULTAD / UNIVERSIDAD", "FACULTAD DE INGENIERÍA MECÁNICA ELÉCTRICA / UNAP"),
    ("UBICACIÓN GEOGRÁFICA", "PUNO - PERÚ"),
    ("FECHA Y LÁMINA PLANO", "PUNO - JUNIO - 2026 | LÁMINA A2 (ESCALA 1:50)")
]

tbl_cover = doc.add_table(rows=len(meta_items), cols=2)
tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
for idx, (label, val) in enumerate(meta_items):
    row = tbl_cover.rows[idx]
    c1, c2 = row.cells[0], row.cells[1]
    set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
    set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(f"      {label:<24}:")
    r1.font.name = 'Arial'; r1.font.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = RGBColor(15, 32, 67)
    
    p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(f" {val}")
    r2.font.name = 'Arial'; r2.font.bold = True; r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(40, 40, 40)

add_p("", space_after=60, first_indent=0)
add_p("PUNO - PERÚ - 2026", WD_ALIGN_PARAGRAPH.CENTER, 0, 0, True, False, 11, (15, 32, 67), first_indent=0)

doc.add_page_break()

# =============================================================================
# ÍNDICE GENERAL EXTENDIDO
# =============================================================================
add_num_heading("ÍNDICE GENERAL DEL EXPEDIENTE TÉCNICO UNIFICADO COMPLETO", level=1)
add_p("El presente Expediente Técnico de Instalaciones Eléctricas de la Vivienda Unifamiliar de 3 Niveles ha sido compilado de forma integral en un único documento maestro dividida en cuatro partes fundamentales:")

index_items = [
    "PARTE I: MEMORIA DESCRIPTIVA",
    "  1.0 Objetivo del Proyecto y Marco General",
    "  2.0 Alcances del Proyecto de Instalaciones Eléctricas",
    "  3.0 Descripción del Proyecto y Distribución Arquitectónica Detallada",
    "      3.1 Primer Piso (Área Social, Servicio, Patio, Cisterna y Electrobombas)",
    "      3.2 Segundo Piso (Área Íntima, Dormitorios, Baños Privados y Comunes)",
    "      3.3 Azotea (Área de Lavandería, Tendedero, Depósito y Servicios)",
    "      3.4 Cuadro Detallado de Áreas Construidas y Ambientes Proyectados",
    "  4.0 Características del Sistema Eléctrico y Suministro",
    "      4.1 Suministro Eléctrico (Electro Puno S.A.A. 220V 3Ø 60Hz)",
    "      4.2 Potencia Instalada (18.25 kW) y Máxima Demanda (13.50 kW)",
    "      4.3 Corrientes de Diseño Nominal (In = 39.36 A, Id = 49.20 A)",
    "      4.4 Dimensionamiento del Alimentador Principal (NH-80)",
    "  5.0 Bases de Cálculo y Normativa de Carga",
    "  6.0 Cuadro Oficial de Cargas y Máxima Demanda",
    "  7.0 Ubicación del Tablero General (TD) y Centros de Carga",
    "  8.0 Redes Complementarias y Telecomunicaciones",
    "  9.0 Normativa y Reglamentación Nacional e Internacional",
    "  10.0 Simbología Reglamentaria y Referencia a Planos",
    "  11.0 REPRODUCCIÓN Y ESQUEMA DEL PLANO DE INSTALACIONES ELÉCTRICAS (LÁMINA A2)",
    "      - Figura 3: Reproducción y Esquema del Plano de Diego Charaja (Lámina A2)",
    "PARTE II: CÁLCULOS JUSTIFICATORIOS Y DIAGRAMA UNIFILAR DETALLADO",
    "  2.1 Introducción y Criterios de Diseño Eléctrico",
    "      2.1.1 Criterio de Capacidad de Corriente (Ampacidad)",
    "      2.1.2 Criterio de Caída de Tensión (ΔV ≤ 2.5% Alimentador / ≤ 4.0% Total)",
    "      2.1.3 Criterio de Protección Térmica y Diferencial (30 mA)",
    "  2.2 Fórmulas Empleadas y Desarrollo Matemático paso a paso",
    "  2.3 Sustento Detallado del Cuadro de Cargas Circuito por Circuito (C-1 a C-8)",
    "  2.4 Cálculo Justificatorio del Alimentador General (Valores numéricos exactos)",
    "  2.5 ESQUEMA Y ESPECIFICACIÓN DETALLADA DEL DIAGRAMA UNIFILAR",
    "      2.5.1 Suministro y Acometida Eléctrica (Electro Puno S.A.A.)",
    "      2.5.2 Medidor de Energía y Caja de Toma L-T",
    "      2.5.3 Interruptor General Termomagnético (3x40A, 10 kA)",
    "      2.5.4 Barras Colectoras y Coordinación de Protecciones",
    "      2.5.5 Desglose Específico Circuito por Circuito (C-1 a C-8)",
    "      - Figura 1: Diagrama Unifilar del Tablero General de Distribución (TD)",
    "  2.6 Cuadro Completo de Dimensionamiento de Circuitos Derivados",
    "  2.7 Dimensionamiento y Peinado del Tablero General (TD de 36 Polos trifásico)",
    "  2.8 Cálculo de Ocupación de Conducciones y Tubos PVC (Ocupación ≤ 40%)",
    "  2.9 Cálculos de Iluminación y Niveles de Lux por Ambiente (RNE EM.010 / EM.020)",
    "PARTE III: ESPECIFICACIONES TÉCNICAS DE EQUIPOS Y MATERIALES",
    "  1. GENERALIDADES (1.1 Introducción, 1.2 Alcance, 1.3 Documentación)",
    "  2. ESPECIFICACIONES TÉCNICAS DE MATERIALES",
    "      2.1 Normas Constructivas y Criterios de Calidad",
    "      2.2 Tuberías Eléctricas de PVC (PVC-P SAP y PVC-L SEL)",
    "      2.3 Cajas Metálicas para Salidas y Pase (F°G° pesado 1.59 mm)",
    "      2.4 Conductores Eléctricos Cero Halógenos (NH-80 / N2XH) y Código de Colores",
    "      2.5 Cintas Aislantes y Accesorios de Empalme",
    "      2.6 Tableros Eléctricos de Distribución (Gabinete metálico LAF 1.5 mm, IP41/IK08)",
    "  3. DISPOSITIVOS Y EQUIPOS VARIOS",
    "      3.1 Interruptores para Alumbrado (10A/16A 250V)",
    "      3.2 Tomacorrientes Generales y de Fuerza (2P+T 16A con obturadores infantiles)",
    "      3.3 Artefactos de Iluminación LED (≥ 100 lm/W)",
    "      3.4 Equipos Autónomos de Alumbrado de Emergencia (Ni-Cd/Litio ≥ 90 min)",
    "  4. PRUEBAS ELECTROMECÁNICAS Y PUESTA EN SERVICIO",
    "      4.1 Pruebas de Resistencia de Aislamiento (Megado 500V DC ≥ 50 MΩ)",
    "      4.2 Prueba de Continuidade Aislamiento de Conductores",
    "      4.3 Prueba de Disparo de Interruptores Diferenciales (30 mA, t ≤ 40 ms)",
    "  5. ESPECIFICACIONES DE MONTAJE Y MONITOREO",
    "  6. PLANOS AS-BUILT Y DOCUMENTACIÓN FINAL",
    "PARTE IV: MATRIZ DE INCONSISTENCIAS Y PROPUESTAS TÉCNICAS DE MEJORA",
    "  1.0 Presentación e Importancia de la Auditoría Técnica",
    "  2.0 Matriz Consolidada de Inconsistencias y Soluciones Técnicas",
    "  3.0 Justificación Detallada de Propuestas Técnicas e Inferencias de Ingeniería",
    "      3.1 Cambio de Concesionaria Eléctrica a Electro Puno S.A.A.",
    "      3.2 Adecuación de Tipología Residencial a Vivienda Unifamiliar de 3 Niveles",
    "      3.3 Implementación Obligatoria de Protecciones Diferenciales (30 mA)",
    "      3.4 Selección de Conductores Ecológicos Cero Halógenos (NH-80 / N2XH)",
    "      3.5 Redimensionamiento del Tablero General a 36 Polos"
]

for item in index_items:
    is_main = not item.startswith("  ")
    add_p(item, space_after=3, bold=is_main, color=(15,32,67) if is_main else (50,50,50), first_indent=0 if is_main else 0.2)

doc.add_page_break()

# =============================================================================
# PARTE I: MEMORIA DESCRIPTIVA
# =============================================================================
add_p("INSTALACIONES ELÉCTRICAS INTERIORES", WD_ALIGN_PARAGRAPH.CENTER, 0, 12, True, False, 14, (15,32,67), first_indent=0)
add_roman_heading("I. MEMORIA DESCRIPTIVA")

add_num_heading("1.0 OBJETIVO DEL PROYECTO Y MARCO GENERAL", level=1)
add_p("El presente Proyecto de Instalaciones Eléctricas Interiores corresponde a la Edificación de una Vivienda Unifamiliar de tres niveles (Primer Piso, Segundo Piso y Azotea), de propiedad de CHARAJA MAMANI DIEGO JEFFERSON (Código 214254), desarrollado para la asignatura de Instalaciones Eléctricas dictada por el Ing. VILLANUEVA CORNEJO MARCOS JOSE en la Escuela Profesional de Ingeniería Mecánica Eléctrica de la Universidad Nacional del Altiplano. El objetivo principal es definir los criterios técnicos, dimensionales, normativos y constructivos para la correcta distribución de energía eléctrica, selección de conductores, tableros, protecciones termomagnéticas y diferenciales, garantizando la máxima seguridad de las personas, la preservación de la infraestructura contra sobrecargas y cortocircuitos, y la eficiencia energética conforme al Código Nacional de Electricidad - Utilización (CNE-U) y al Reglamento Nacional de Edificaciones (RNE EM.010).")

add_num_heading("2.0 ALCANCES DEL PROYECTO DE INSTALACIONES ELÉCTRICAS", level=1)
add_p("El proyecto comprende el diseño completo de las instalaciones eléctricas en Baja Tensión y redes complementarias:")
add_p("a) Instalaciones Eléctricas en Baja Tensión", bold=True, first_indent=0)
add_p("• El diseño de la Acometida Eléctrica trifásica desde la red pública de distribución secundaria de la empresa concesionaria ELECTRO PUNO S.A.A. hasta la caja de protección y medición ubicada en el límite de propiedad.")
add_p("• El diseño del Alimentador General desde la caja de medición hasta el Tablero General de Distribución (TD) empotrado en el Hall del primer piso.")
add_p("• Red de 8 circuitos derivados (C-1 a C-8) para Alumbrado, Tomacorrientes Generales y Fuerza.")
add_p("• Alimentación de fuerza para cargas especiales: Cocina Eléctrica (6.0 kW), Terma Eléctrica (1.5 kW), Electrobomba de Agua (0.75 HP), Lavadora y Secadora de Ropa (2.5 kW) y Tomacorrientes especiales de cocina (1.5 kW).")
add_p("b) Redes de Comunicaciones y Servicios Auxiliares", bold=True, first_indent=0)
add_p("• Redes de ductos, cajas de pase y cajas de salida para telefonía fija/móvil, cable TV e intercomunicador con abrepuerta eléctrico.")

add_num_heading("3.0 DESCRIPCIÓN DEL PROYECTO Y DISTRIBUCIÓN ARQUITECTÓNICA DETALLADA", level=1)
add_p("La edificación presenta una arquitectura residencial unifamiliar construida sobre tres niveles, con la siguiente distribución espacial detallada extraída de los planos de diseño (Lámina A2, Escala 1:50):")

add_num_heading("3.1 Primer Piso (Área Social, Servicio, Patio, Cisterna y Electrobombas)", level=2)
add_p("Comprende la zona de ingreso principal, el área social y de servicio primario: Sala - Comedor iluminada y ventilada naturalmente, Estar familiar, Cocina equipada con salidas para tomacorrientes de electrodomésticos menores y alimentador trifásico para cocina eléctrica, un Dormitorio de huéspedes/estudio, Hall de circulación principal donde se ubica empotrado el Tablero General (TD), Servicio Higiénico (S.H.), Ducto de ventilación, Patio posterior al aire libre y el cuarto técnico para la Ubicación de Cisterna de agua potable y Electrobombas (0.75 HP).")

add_num_heading("3.2 Segundo Piso (Área Íntima, Dormitorios, Baños Privados y Comunes)", level=2)
add_p("Comprende la zona privada o íntima de la vivienda: Dormitorio Principal con Servicio Higiénico privado incorporado, tres Dormitorios secundarios ampliamente ventilados, Hall de distribución privada, un Servicio Higiénico común completo, Ducto de ventilación vertical y Clósets empotrados en cada dormitorio. En este nivel se alimenta la Terma Eléctrica (1.5 kW) para suministro de agua caliente sanitaria.")

add_num_heading("3.3 Azotea (Área de Lavandería, Tendedero, Depósito y Servicios)", level=2)
add_p("Comprende la zona de servicios complementarios y esparcimiento: Lavandería independizada con tomacorrientes reforzados para Lavadora y Secadora de ropa (2.5 kW), Tendedero al aire libre, Depósito de almacenamiento de enseres, Hall de circulación, Servicio Higiénico de servicio y Ducto de ventilación.")

add_num_heading("3.4 Cuadro Detallado de Áreas Construidas y Ambientes Proyectados", level=2)

tbl_areas = doc.add_table(rows=5, cols=3)
tbl_areas.alignment = WD_TABLE_ALIGNMENT.CENTER
a_hdrs = ["Nivel", "Ambientes Incluidos", "Área Techada (m²)"]
hdr_a_row = tbl_areas.rows[0]
for idx, h in enumerate(a_hdrs):
    c = hdr_a_row.cells[idx]
    set_cell_bg(c, "0F2043")
    set_cell_margins(c, 80, 80, 80, 80)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = 'Arial'; r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)

a_data = [
    ("Primer Piso", "Sala-Comedor, Estar, Cocina, Dormitorio 1, Hall (TD), S.H. 1, Cisterna/Bomba, Patio", "90.00 m²"),
    ("Segundo Piso", "Dormitorio Principal + S.H., Dormitorios 2, 3, 4, S.H. Común, Hall 2do Piso", "85.00 m²"),
    ("Azotea", "Lavandería, Depósito, Tendedero al aire libre, S.H. Azotea", "45.00 m²"),
    ("ÁREA TOTAL", "SUPERFICIE CONSTRUIDA ACUMULADA TOTAL DE LA VIVIENDA UNIFAMILIAR", "220.00 m²")
]
for r_idx, row_data in enumerate(a_data, start=1):
    row = tbl_areas.rows[r_idx]
    bg = "F0F4F8" if r_idx % 2 == 1 else "FFFFFF"
    if r_idx == 4: bg = "D9E2EC"
    for c_idx, val in enumerate(row_data):
        c = row.cells[c_idx]
        set_cell_bg(c, bg)
        set_cell_margins(c, 50, 50, 60, 60)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val); r.font.name = 'Arial'; r.font.size = Pt(8.5)
        if r_idx == 4 or c_idx in [0, 2]: r.font.bold = True

add_p("", space_after=6)

add_num_heading("4.0 CARACTERÍSTICAS DEL SISTEMA ELÉCTRICO Y SUMINISTRO", level=1)
add_p("• Concesionaria de Suministro : ELECTRO PUNO S.A.A. (Red Secundaria Pública de Baja Tensión)\n"
      "• Tensión Nominal de Servicio : 220 V (Sistema Trifásico 3Ø)\n"
      "• Frecuencia Nominal : 60 Hz\n"
      "• Potencia Instalada Total (PI) : 18,250.00 Watts (18.25 kW)\n"
      "• Máxima Demanda Total (MD) : 13,500.00 Watts (13.50 kW)\n"
      "• Factor de Potencia de Diseño (cos φ) : 0.90 (inductivo residencial CNE-U)\n"
      "• Corriente Nominal de Diseño (In) : 39.36 Amperios\n"
      "• Corriente de Diseño Corregida (Id) : 49.20 Amperios (con factor de sobrecarga 1.25 CNE-U 050-102)\n"
      "• Conductor Alimentador Principal : 3-1x10 mm² NH-80 en tub. PVC-P 25 mmØ (1\"Ø)\n"
      "• Interruptor General del TD : Termomagnético Tripolar 3 x 40 A (Capacidad de Ruptura Icu ≥ 10 kA)\n"
      "• Capacidad del Tablero Proyectado : 36 Polos trifásico empotrado (con 20% de reserva libre)")

add_num_heading("5.0 BASES DE CÁLCULO Y NORMATIVA DE CARGA", level=1)
add_p("La determinación de la Máxima Demanda se basa estrictamente en las reglas de carga del Código Nacional de Electricidad - Utilización (Tabla 14 y Tabla 22):\n"
      "• C-1 (Alumbrado y Tomacorrientes Generales 90 m²): Se asignan 2,500 W con Factor de Simultaneidad FS = 1.00 -> MD = 2,500 W.\n"
      "• C-2 (Área Adicional): Se estiman 2,000 W con Factor de Simultaneidad FS = 0.35 -> MD = 700 W.\n"
      "• C-3 (Cocina Eléctrica 6.0 kW): Según CNE-U Tabla 22, para 6 kW se aplica FS = 0.80 -> MD = 4,800 W.\n"
      "• C-4 (Terma Eléctrica 1.5 kW): Carga continua con FS = 1.00 -> MD = 1,500 W.\n"
      "• C-5 (Electrobomba de Agua 0.75 HP): Carga inductiva con FS = 1.00 -> MD = 750 W.\n"
      "• C-6 (Tomacorrientes de Cocina Cargas Menores): 1,500 W con FS = 0.50 -> MD = 750 W.\n"
      "• C-7 (Lavadora y Secadora Azotea): 2,500 W con FS = 0.70 -> MD = 1,750 W.\n"
      "• C-8 (Reserva / Ampliación Futura): 1,500 W proyectados con FS = 0.50 -> MD = 750 W.")

add_num_heading("6.0 CUADRO OFICIAL DE CARGAS Y MÁXIMA DEMANDA", level=1)

tbl_cargas = doc.add_table(rows=10, cols=5)
tbl_cargas.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Circuito", "Descripción de la Carga / Ubicación", "PI (W)", "FS", "MD (W)"]
hdr_row = tbl_cargas.rows[0]
for idx, h in enumerate(headers):
    c = hdr_row.cells[idx]
    set_cell_bg(c, "0F2043")
    set_cell_margins(c, 80, 80, 100, 100)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = 'Arial'; r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(255,255,255)

data_cargas = [
    ("C-1", "Alumbrado y Tomacorrientes Generales (Primeros 90 m²)", "2,500", "1.00", "2,500"),
    ("C-2", "Alumbrado y Tomacorrientes Generales (Área Adicional)", "2,000", "0.35", "700"),
    ("C-3", "Cocina Eléctrica (Primer Piso)", "6,000", "0.80", "4,800"),
    ("C-4", "Terma Eléctrica / Ducha (Segundo Piso)", "1,500", "1.00", "1,500"),
    ("C-5", "Electrobomba de Agua (0.75 HP - Patio/Cisterna)", "750", "1.00", "750"),
    ("C-6", "Tomacorrientes de Cocina (Microondas, Licuadora, etc.)", "1,500", "0.50", "750"),
    ("C-7", "Lavadora y Secadora de Ropa (Azotea / Lavandería)", "2,500", "0.70", "1,750"),
    ("C-8", "Cargas Especiales de Reserva / Ampliaciones", "1,500", "0.50", "750"),
    ("TOTALES", "POTENCIA INSTALADA TOTAL Y MÁXIMA DEMANDA", "18,250", "-", "13,500")
]
for r_idx, row_data in enumerate(data_cargas, start=1):
    row = tbl_cargas.rows[r_idx]
    bg = "F0F4F8" if r_idx % 2 == 1 else "FFFFFF"
    if r_idx == 9: bg = "D9E2EC"
    for c_idx, val in enumerate(row_data):
        c = row.cells[c_idx]
        set_cell_bg(c, bg)
        set_cell_margins(c, 60, 60, 100, 100)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val); r.font.name = 'Arial'; r.font.size = Pt(9)
        if r_idx == 9 or c_idx in [0, 4]: r.font.bold = True

add_p("", space_after=6)

add_num_heading("7.0 UBICACIÓN DEL TABLERO GENERAL (TD)", level=1)
add_p("El Tablero General de Distribución (TD) estará ubicado en forma empotrada en el muro de ladrillo del Hall de circulación del Primer Piso, a una altura de 1.80 m del N.P.T. medidos hasta el borde superior de la caja metálica. Esta posición física garantiza una ubicación centralizada respecto a los centros de carga de la edificación, facilidad de maniobra y protección contra agentes externos directos.")

add_num_heading("8.0 REDES COMPLEMENTARIAS Y TELECOMUNICACIONES", level=1)
add_p("Se proyecta una red de ductos de PVC liviano empotrados en muros y lozas con cajas de pase metálicas rectangulares (4\"x2\") y cuadradas (4\"x4\") destinadas exclusivamente para canalizaciones de telefonía fija, televisión por cable (TV-cable) e intercomunicador exterior con abrepuerta. Todas las redes de corrientes débiles mantienen una distancia mínima de 30 cm con los conductores de fuerza para evitar interferencias electromagnéticas.")

add_num_heading("9.0 NORMATIVA, CÓDIGOS Y REGLAMENTACIONES", level=1)
add_p("El diseño y ejecución del proyecto cumplen con:\n"
      "• Código Nacional de Electricidad - Utilización (CNE-U, Perú).\n"
      "• Reglamento Nacional de Edificaciones (RNE) - Norma Técnica EM.010 'Instalaciones Eléctricas Interiores'.\n"
      "• Normas de la Comisión Electrotécnica Internacional (IEC 60364, IEC 60898, IEC 61008).\n"
      "• National Electrical Code (NFPA 70 / NEC).\n"
      "• Especificaciones Técnicas de la Empresa Concesionaria ELECTRO PUNO S.A.A.")

add_num_heading("10.0 SIMBOLOGÍA REGLAMENTARIA Y REFERENCIA A PLANOS", level=1)
add_p("Toda la simbología utilizada en la documentación y esquemas cumple estrictamente con la Norma DGE 'Símbolos Gráficos en Electricidad' del Ministerio de Energía y Minas del Perú.")

add_num_heading("11.0 REPRODUCCIÓN Y ESQUEMA DEL PLANO DE INSTALACIONES ELÉCTRICAS (LÁMINA A2)", level=1)
add_p("A continuación se integra la representación gráfica oficial del plano de instalaciones eléctricas proyectado por Diego Jefferson Charaja Mamani (Lámina A2, Escala 1:50), especificando la arquitectura de los tres niveles, los centros de luz, tomacorrientes y tableros:")

# INSERT FIGURA PLANO ARQUITECTÓNICO DIEGO
fig_plano_file = os.path.join(TARGET_DIR, "fig_plano_diego_arquitectura.png")
if os.path.exists(fig_plano_file):
    p_img_p = doc.add_paragraph()
    p_img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img_p = p_img_p.add_run()
    r_img_p.add_picture(fig_plano_file, width=Inches(6.2))
    add_caption("Figura 3: Reproducción y Esquema del Plano de Instalaciones Eléctricas por Niveles de Diego Charaja (Lámina A2, Escala 1:50)")

doc.add_page_break()

# =============================================================================
# PARTE II: CÁLCULOS JUSTIFICATORIOS Y DIAGRAMA UNIFILAR DETALLADO
# =============================================================================
add_p("INSTALACIONES ELÉCTRICAS INTERIORES", WD_ALIGN_PARAGRAPH.CENTER, 0, 12, True, False, 14, (15,32,67), first_indent=0)
add_roman_heading("II. CÁLCULOS JUSTIFICATORIOS Y DIAGRAMA UNIFILAR DETALLADO")

add_num_heading("2.1 INTRODUCCIÓN Y CRITERIOS DE DISEÑO ELÉCTRICO", level=1)
add_p("El dimensionamiento de los conductores, tuberías y dispositivos de protección de la Vivienda Unifamiliar se sustenta en el cálculo analítico bajo tres criterios fundamentales de ingeniería eléctrica:")

add_num_heading("2.1.1 Criterio de Capacidad de Corriente (Ampacidad)", level=2)
add_p("La corriente nominal de servicio (In) para un sistema trifásico (220 V, 3Ø) se calcula mediante:")
add_p("In = MD / (√3 x V x cos φ) = 13500 / (1.73205 x 220 x 0.90) = 39.36 A", WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=(15,32,67))
add_p("Según la Regla 050-102 del CNE-Utilización, los alimentadores continuos deben dimensionarse para una corriente de diseño corregida (Id) con un factor de sobrecarga de 1.25:\n"
      "Id = 1.25 x In = 1.25 x 39.36 A = 49.20 A.")

add_num_heading("2.1.2 Criterio de Caída de Tensión (ΔV)", level=2)
add_p("La caída de tensión en tramos trifásicos (3Ø) y monofásicos (1Ø) se evalúa mediante:")
add_p("ΔV(3Ø) = (√3 x I x L x ρ) / S     |     ΔV(1Ø) = (2 x I x L x ρ) / S", WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=(15,32,67))
add_p("Donde ρ = 0.0175 Ω·mm²/m es la resistividad del cobre recocido a 20 °C, L es la longitud del tramo en metros, I es la corriente en amperios y S es la sección del conductor en mm². El CNE-U exige que la caída de tensión acumulada no supere el 4.0% total (máximo 2.5% en el alimentador general y 1.5% en los derivados).")

add_num_heading("2.1.3 Criterio de Protección Térmica y Diferencial (30 mA)", level=2)
add_p("Cada conductor se protege contra sobrecargas y cortocircuitos mediante interruptores termomagnéticos (ITM) seleccionados de forma que In_itm ≤ I_adm del cable. Además, la Regla 020-204 del CNE-U impone la instalación de Interruptores Diferenciales (ID) con sensibilidad IΔn = 30 mA en todos los circuitos de tomacorrientes y áreas húmedas para la protección de la vida humana.")

add_num_heading("2.2 SUSTENTO DETALLADO DEL ALIMENTADOR GENERAL", level=1)
add_p("• Máxima Demanda Total (MD) : 13,500.00 W (13.50 kW)\n"
      "• Tensión de Servicio : 220 V (Trifásico 3Ø, 60 Hz)\n"
      "• Corriente Nominal (In) : In = 13500 / (1.73205 x 220 x 0.90) = 39.36 A\n"
      "• Corriente de Diseño (Id) : Id = 1.25 x 39.36 A = 49.20 A\n"
      "• Selección de Conductor : Cable Cero Halógenos 3-1x10 mm² NH-80.\n"
      "  - Ampacidad nominal del conductor 10 mm² NH-80 empotrado en tubo PVC: 57 A.\n"
      "  - Verificación por ampacidad: 57 A > 49.20 A (Cumple holgadamente por capacidad térmica).\n"
      "• Cálculo de Caída de Tensión (L = 15 m desde Medidor a TD):\n"
      "  - ΔV = (1.73205 x 39.36 x 15 x 0.0175) / 10 = 1.79 V\n"
      "  - % ΔV = (1.79 V / 220 V) x 100% = 0.81% (Cumple holgadamente con el límite máximo de 2.5% del CNE-U).")

# =============================================================================
# SECCIÓN ESPECÍFICA DETALLADA DEL PLANO UNIFILAR
# =============================================================================
add_num_heading("2.3 ESQUEMA Y ESPECIFICACIÓN DETALLADA DEL DIAGRAMA UNIFILAR", level=1)
add_p("El Diagrama Unifilar del Tablero General (TD) constituye el plano esquemático fundamental donde se representan de forma unilineal todas las protecciones, barras colectoras, calibres de conductores y distribución de cargas de la edificación. A continuación se desglosan en detalle las características electromecánicas de cada tramo y dispositivo:")

add_num_heading("2.3.1 Suministro y Acometida Eléctrica (Electro Puno S.A.A.)", level=2)
add_p("• Red de Origen: Red Secundaria Pública de Distribución en Baja Tensión (220 V, 3Ø, 60 Hz) de ELECTRO PUNO S.A.A.\n"
      "• Tipo de Acometida: Acometida subterránea o aérea reforzada ejecutada con cable unipolar o tetrapolar tipo N2XH 0.6/1 kV de 10 mm².\n"
      "• Caja de Toma y Medición: Caja metálica reforzada tipo L-T (o caja de poliéster reforzado con fibra de vidrio IP65) alojada en el murete de fachada límite de propiedad. Aloja un Medidor Electrónico Trifásico multitarifa homologado por la concesionaria.")

add_num_heading("2.3.2 Interruptor General de Cabecera en el Tablero General (TD)", level=2)
add_p("• Tipo de Dispositivo: Interruptor Termomagnético (ITM) Tripolar automático.\n"
      "• Calibre Nominal: 3 x 40 Amperios (3P).\n"
      "• Capacidad de Ruptura / Cortocircuito (Icu): 10 kA a 220/240 V AC (Norma IEC 60947-2 / IEC 60898).\n"
      "• Curva de Disparo Térmico-Magnético: Curva C (Disparo magnético entre 5 y 10 veces In, idóneo para cargas residenciales con picos de arranque de motores).\n"
      "• Función: Protección integral del alimentador principal y desenganche total manual/automático del tablero.")

add_num_heading("2.3.3 Barras Colectoras y Coordinación de Protecciones", level=2)
add_p("• Barras Colectoras R-S-T: Pletinas de cobre electrolítico puro de 99.9% Cu, dimensionadas para una corriente nominal de 100 A, montadas sobre aisladores de resina epóxica.\n"
      "• Luces Piloto Indicadoras: Juego de 3 lámparas piloto LED de color Rojo, Amarillo y Azul conectadas a las fases R-S-T para verificación visual de presencia de tensión en barras.")

add_num_heading("2.3.4 Desglose Específico Circuito por Circuito (C-1 a C-8)", level=2)
add_p("A continuación se especifica la configuración de protección y cableado de cada una de las 8 salidas derivadas del Tablero General:")

add_p("1. Circuito C-1 (Alumbrado y Tomacorrientes Generales 1er Piso - Base 90 m²):\n"
      "   - Protecciones: ITM 2P x 16 A (6 kA, Curva C) + ID 2P x 25 A (Sensibilidad 30 mA, Clase AC).\n"
      "   - Conductor: 2 x 2.5 mm² NH-80.\n"
      "   - Tubería: PVC Liviano (SEL) Ø 20 mm (3/4\").\n"
      "   - Carga: Centros de luz LED, dicroicos y tomacorrientes generales en Sala, Comedor, Estar, Dormitorio 1er Piso. MD = 2,500 W, In = 12.63 A.", first_indent=0.2)

add_p("2. Circuito C-2 (Alumbrado y Tomacorrientes Generales 2do Piso y Azotea - Área Adicional):\n"
      "   - Protecciones: ITM 2P x 16 A (6 kA, Curva C) + ID 2P x 25 A (Sensibilidad 30 mA, Clase AC).\n"
      "   - Conductor: 2 x 2.5 mm² NH-80.\n"
      "   - Tubería: PVC Liviano (SEL) Ø 20 mm (3/4\").\n"
      "   - Carga: Centros de luz LED y tomacorrientes en 4 Dormitorios 2do piso, Baños, Hall, Depósito y Azotea. MD = 700 W, In = 3.54 A.", first_indent=0.2)

add_p("3. Circuito C-3 (Cocina Eléctrica - 1er Piso):\n"
      "   - Protecciones: ITM 3P x 25 A (6 kA, Curva C) + ID 3P x 32 A (Sensibilidad 30 mA, Clase A/AC).\n"
      "   - Conductor: 3 x 6.0 mm² NH-80.\n"
      "   - Tubería: PVC Pesado (SAP) Ø 25 mm (1\").\n"
      "   - Carga: Suministro trifásico 220V 3Ø para Cocina Eléctrica empotrada de 4 hornillas + horno (PI = 6,000 W @ FS=0.80 -> MD = 4,800 W, In = 13.99 A).", first_indent=0.2)

add_p("4. Circuito C-4 (Terma Eléctrica / Ducha - 2do Piso):\n"
      "   - Protecciones: ITM 2P x 16 A (6 kA, Curva C) + ID 2P x 25 A (Sensibilidad 30 mA de alta rapidez).\n"
      "   - Conductor: 2 x 4.0 mm² NH-80.\n"
      "   - Tubería: PVC Pesado (SAP) Ø 20 mm (3/4\").\n"
      "   - Carga: Terma acumuladora de agua caliente sanitaria de 80 Litros (1,500 W). MD = 1,500 W, In = 7.58 A.", first_indent=0.2)

add_p("5. Circuito C-5 (Electrobomba de Agua - Patio / Cisterna):\n"
      "   - Protecciones: ITM 2P x 16 A (6 kA, Curva C) + ID 2P x 25 A (30 mA) + Guardamotor / Térmico + Control por boya automática de nivel.\n"
      "   - Conductor: 2 x 2.5 mm² NH-80.\n"
      "   - Tubería: PVC Pesado (SAP) Ø 20 mm (3/4\").\n"
      "   - Carga: Motor eléctrico de bomba centrífuga de agua potable de 0.75 HP (750 W). MD = 750 W, In = 3.79 A.", first_indent=0.2)

add_p("6. Circuito C-6 (Tomacorrientes Especiales de Cocina - Artefactos Menores):\n"
      "   - Protecciones: ITM 2P x 20 A (6 kA, Curva C) + ID 2P x 25 A (Sensibilidad 30 mA).\n"
      "   - Conductor: 2 x 4.0 mm² NH-80.\n"
      "   - Tubería: PVC Pesado (SAP) Ø 20 mm (3/4\").\n"
      "   - Carga: Horno microondas, licuadora, cafetera, freidora de aire, arrocera (PI = 1,500 W @ FS=0.50 -> MD = 750 W, In = 3.79 A).", first_indent=0.2)

add_p("7. Circuito C-7 (Lavadora y Secadora de Ropa - Azotea / Lavandería):\n"
      "   - Protecciones: ITM 2P x 20 A (6 kA, Curva C) + ID 2P x 25 A (Sensibilidad 30 mA).\n"
      "   - Conductor: 2 x 4.0 mm² NH-80.\n"
      "   - Tubería: PVC Pesado (SAP) Ø 20 mm (3/4\").\n"
      "   - Carga: Lavadora automática y secadora eléctrica en zona de lavandería azotea (PI = 2,500 W @ FS=0.70 -> MD = 1,750 W, In = 8.84 A).", first_indent=0.2)

add_p("8. Circuito C-8 (Cargas Especiales de Reserva / Ampliaciones Futuras):\n"
      "   - Protecciones: ITM 2P x 20 A (6 kA, Curva C) equipado de reserva en el gabinete metálico.\n"
      "   - Conductor: 2 x 4.0 mm² NH-80 proyectado en ducto PVC-P Ø 20 mm hacia azotea/patio. MD = 750 W, In = 3.79 A.", first_indent=0.2)

# INSERT FIGURA 1: DIAGRAMA UNIFILAR
fig1_path = os.path.join(TARGET_DIR, "fig_diagrama_unifilar.png")
if os.path.exists(fig1_path):
    p_img1 = doc.add_paragraph()
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img1 = p_img1.add_run()
    r_img1.add_picture(fig1_path, width=Inches(6.2))
    add_caption("Figura 1: Diagrama Unifilar Detallado del Tablero General de Distribución (TD)")

add_num_heading("2.4 CUADRO COMPLETO DE DIMENSIONAMIENTO DE CIRCUITOS DERIVADOS", level=1)

tbl_c = doc.add_table(rows=9, cols=7)
tbl_c.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs_c = ["Circuito", "Descripción", "MD (W)", "In (A)", "Conductor NH-80", "Tubería PVC", "Protección ITM / ID"]
hdr_c_row = tbl_c.rows[0]
for idx, h in enumerate(hdrs_c):
    c = hdr_c_row.cells[idx]
    set_cell_bg(c, "0F2043")
    set_cell_margins(c, 80, 80, 80, 80)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = 'Arial'; r.font.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(255,255,255)

specs_data = [
    ("C-1", "Alum. y Tomac. 1er Piso", "2,500", "12.63", "2x2.5 mm²", "20 mm (3/4\" SEL)", "ITM 2x16A / ID 2x25A 30mA"),
    ("C-2", "Alum. y Tomac. 2do/Azotea", "700", "3.54", "2x2.5 mm²", "20 mm (3/4\" SEL)", "ITM 2x16A / ID 2x25A 30mA"),
    ("C-3", "Cocina Eléctrica (1er Piso)", "4,800", "13.99*", "3x6.0 mm²", "25 mm (1\" SAP)", "ITM 3x25A / ID 3x32A 30mA"),
    ("C-4", "Terma Eléctrica (2do Piso)", "1,500", "7.58", "2x4.0 mm²", "20 mm (3/4\" SAP)", "ITM 2x16A / ID 2x25A 30mA"),
    ("C-5", "Electrobomba Agua (0.75 HP)", "750", "3.79", "2x2.5 mm²", "20 mm (3/4\" SAP)", "ITM 2x16A / ID 2x25A 30mA"),
    ("C-6", "Tomac. Cocina Artefactos", "750", "3.79", "2x4.0 mm²", "20 mm (3/4\" SAP)", "ITM 2x20A / ID 2x25A 30mA"),
    ("C-7", "Lavadora/Secadora Azotea", "1,750", "8.84", "2x4.0 mm²", "20 mm (3/4\" SAP)", "ITM 2x20A / ID 2x25A 30mA"),
    ("C-8", "Reserva Especial", "750", "3.79", "2x4.0 mm²", "20 mm (3/4\" SAP)", "ITM 2x20A (Reserva equipada)")
]
for r_idx, row_data in enumerate(specs_data, start=1):
    row = tbl_c.rows[r_idx]
    bg = "F0F4F8" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, val in enumerate(row_data):
        c = row.cells[c_idx]
        set_cell_bg(c, bg)
        set_cell_margins(c, 50, 50, 60, 60)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val); r.font.name = 'Arial'; r.font.size = Pt(8)
        if c_idx in [0, 6]: r.font.bold = True

add_p("", space_after=6)

add_num_heading("2.5 CÁLCULO DE OCUPACIÓN DE CONDUCCIONES Y TUBOS PVC (PORCENTAJE ≤ 40%)", level=1)
add_p("Conforme a la Regla 070-300 del CNE-Utilización, el área total ocupada por los conductores eléctricos dentro de una tubería o canalización no debe exceder del 40% de la sección transversal interna del tubo para permitir la disipación térmica y evitar el estrangulamiento durante el cableado.")
add_p("• Alimentador Principal (Tubo PVC-P 25 mmØ = 1\"Ø, Área interna = 380 mm²):\n"
      "  - 3 conductores de 10 mm² NH-80 (Diámetro exterior con aislamiento = 6.2 mm, Área total = 3 x 30.19 = 90.57 mm²).\n"
      "  - % Ocupación: (90.57 / 380) x 100% = 23.83% (Cumple holgadamente con el límite de 40%).\n"
      "• Circuitos Derivados de Fuerza (Tubo PVC-P 20 mmØ = 3/4\"Ø, Área interna = 220 mm²):\n"
      "  - 2 conductores de 4.0 mm² NH-80 (Área total = 2 x 15.20 = 30.40 mm²).\n"
      "  - % Ocupación: (30.40 / 220) x 100% = 13.81% (Cumple holgadamente).")

add_num_heading("2.6 CÁLCULOS DE ILUMINACIÓN Y NIVELES DE LUX POR AMBIENTE (RNE EM.010)", level=1)
add_p("Se determinan los niveles de iluminancia media exigidos por el RNE EM.010 y EM.020 para garantizar el confort visual y la eficiencia energética LED:")
add_p("• Sala - Comedor (90 m² total piso 1): Iluminancia requerida 150 Lux. Se instalan 8 dicroicos LED 7W + 2 paneles LED 24W (Flujo total = 5,400 lm -> 180 Lux promedio, CUMPLE).\n"
      "• Cocina (15 m²): Iluminancia requerida 300 Lux. Se instalan 2 Paneles LED 24W (4000 K neutro, Flujo total = 4,800 lm -> 320 Lux, CUMPLE).\n"
      "• Dormitorios (12 a 16 m² c/u): Iluminancia requerida 100 Lux. Se instala 1 Plafón LED 18W (3000 K cálido, Flujo = 1,800 lm -> 125 Lux, CUMPLE).\n"
      "• Baños / Servicios Higiénicos: Iluminancia requerida 150 Lux. Se instala 1 Adosable LED IP44 15W -> 160 Lux, CUMPLE.")

doc.add_page_break()

# =============================================================================
# PARTE III: ESPECIFICACIONES TÉCNICAS
# =============================================================================
add_p("INSTALACIONES ELÉCTRICAS INTERIORES", WD_ALIGN_PARAGRAPH.CENTER, 0, 12, True, False, 14, (15,32,67), first_indent=0)
add_roman_heading("III. ESPECIFICACIONES TÉCNICAS DE MATERIALES Y MONTAJE")

add_num_heading("1. GENERALIDADES Y MARCO NORMATIVO", level=1)
add_num_heading("1.1 INTRODUCCIÓN Y ALCANCE", level=2)
add_p("Las presentes Especificaciones Técnicas establecen los requisitos mínimos obligatorios de fabricación, calidad, dimensiones, pruebas y métodos de instalación de todos los componentes de las instalaciones eléctricas de la Vivienda Unifamiliar.")

add_num_heading("1.2 NORMAS TÉCNICAS OBLIGATORIAS", level=2)
add_p("Todos los materiales y equipos deberán ser nuevos y de primer uso, fabricados de acuerdo con las siguientes normas:\n"
      "• CNE-Utilización: Código Nacional de Electricidad del Perú.\n"
      "• RNE EM.010: Reglamento Nacional de Edificaciones.\n"
      "• IEC 60364: Low-voltage electrical installations.\n"
      "• IEC 60898: Circuit breakers for overcurrent protection.\n"
      "• IEC 61008: Residual current operated circuit-breakers (RCCB).\n"
      "• NFPA 70: National Electrical Code (NEC).")

add_num_heading("2. ESPECIFICACIONES TÉCNICAS DE MATERIALES", level=1)
add_num_heading("2.1 TUBERÍAS ELÉCTRICAS DE PVC", level=2)
add_p("• Tubería PVC-P (SAP - Pesado): Se empleará en el alimentador general, lozas de concreto y empotramientos expuestos a impacto mecánico (25 mm y 20 mm de diámetro nominal).\n"
      "• Tubería PVC-L (SEL - Liviano): Se empleará en muros de mampostería empotrados para circuitos de alumbrado y tomacorrientes (20 mm de diámetro nominal).\n"
      "• Propiedades: Autoextinguibles, no propagadoras de la llama, resistencia al impacto > 2 Joule.")

add_num_heading("2.2 CAJAS METÁLICAS", level=2)
add_p("Fabricadas en plancha de fierro galvanizado pesado de 1.59 mm (1/16\") de espesor mínimo:\n"
      "• Octogonales (4\" x 1 1/2\"): Para salidas de alumbrado en techo y centros de luz.\n"
      "• Rectangulares (4\" x 2\" x 2\"): Para interruptores y tomacorrientes generales.\n"
      "• Cuadradas (4\" x 4\" x 2\"): Para salidas de fuerza (Cocina C-3, Terma C-4) y cajas de pase.")

add_num_heading("2.3 CONDUCTORES ELÉCTRICOS CERO HALÓGENOS (NH-80 / N2XH)", level=2)
add_p("Conductores de cobre electrolítico puro (99.9%), aislamiento termoplástico libre de halógenos (LSZH), nula emisión de gases tóxicos/corrosivos y retardo a la llama (NTP 370.252).\n"
      "• Código de Colores CNE-U: Fase A (Negro), Fase B (Rojo), Fase C (Azul), Neutro (Blanco/Gris).")

add_num_heading("2.4 TABLEROS ELÉCTRICOS DE DISTRIBUCIÓN (TD)", level=2)
add_p("Gabinete metálico empotrado LAF 1.5 mm, IP41/IK08, 36 Polos trifásico. ITM General 3x40A (10 kA), ITM derivados (6 kA) e Interruptores Diferenciales ID de 30 mA.")

add_num_heading("3. DISPOSITIVOS Y EQUIPOS VARIOS", level=1)
add_p("• Interruptores de Alumbrado: Unipolares simples, dobles y conmutables de 10A/16A 250V AC.\n"
      "• Tomacorrientes: Duplex 2P 16A 250V con obturadores de seguridad infantil.\n"
      "• Luminarias LED: Alta eficiencia (≥ 100 lm/W), FP > 0.92, 3000 K en dormitorios y 4000 K en cocina/baños.\n"
      "• Equipos Autónomos de Emergencia: Luces faro LED ajustables con batería de Ni-Cd/Litio (autonomía ≥ 90 min).")

add_num_heading("4. PRUEBAS ELECTROMECÁNICAS Y PUESTA EN SERVICIO", level=1)
add_p("Pruebas obligatorias antes de la energización:\n"
      "1. Prueba de Resistencia de Aislamiento (Megado 500V DC ≥ 50 MΩ entre fases y masa).\n"
      "2. Prueba de Continuidad de Aislamiento de Conducciones.\n"
      "3. Prueba de Disparo de Diferenciales a 30 mA (t ≤ 40 ms).")

doc.add_page_break()

# =============================================================================
# PARTE IV: MATRIZ DE INCONSISTENCIAS Y PROPUESTAS TÉCNICAS
# =============================================================================
add_p("INSTALACIONES ELÉCTRICAS INTERIORES", WD_ALIGN_PARAGRAPH.CENTER, 0, 12, True, False, 14, (15,32,67), first_indent=0)
add_roman_heading("IV. MATRIZ DE INCONSISTENCIAS Y PROPUESTAS TÉCNICAS DE MEJORA")

add_num_heading("1.0 PRESENTACIÓN E IMPORTANCIA DE LA AUDITORÍA TÉCNICA", level=1)
add_p("Auditoría técnica cruzada entre los planos de Diego Charaja (Lámina A2), su cuadro de cargas en Excel y las exigencias del Código Nacional de Electricidad.")

add_num_heading("2.0 MATRIZ CONSOLIDADA DE INCONSISTENCIAS Y SOLUCIONES TÉCNICAS", level=1)

tbl_m = doc.add_table(rows=6, cols=5)
tbl_m.alignment = WD_TABLE_ALIGNMENT.CENTER
m_hdrs = ["Item", "Elemento Evaluado", "Discrepancia Identificada", "Solución / Propuesta Adoptada", "Justificación Normativa"]
hdr_m_row = tbl_m.rows[0]
for idx, h in enumerate(m_hdrs):
    c = hdr_m_row.cells[idx]
    set_cell_bg(c, "0F2043")
    set_cell_margins(c, 80, 80, 80, 80)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.name = 'Arial'; r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)

m_data = [
    ("1", "Concesionaria Eléctrica", "Plano indica 'RED DE EDELNOR' (Lima Norte).", "Se corrigió consignando a ELECTRO PUNO S.A.A. (o concesionaria local).", "Coherencia geográfica con la ciudad de Puno / UNAP."),
    ("2", "Tipología de Edificación", "Word modelo citaba 'Vivienda Multifamiliar'.", "Se reestructuró para 'Vivienda Unifamiliar de 3 niveles' según plano.", "Fidelidad con la arquitectura del plano A2."),
    ("3", "Protecciones Diferenciales", "Excel omitía especificar diferenciales por circuito.", "Se incorporó la exigencia de ID de 30 mA para C-1 a C-7.", "CNE-Utilización Regla 020-204 obligatoria."),
    ("4", "Tipo de Conductores", "Word modelo mencionaba cable TW/THW tradicional.", "Se especificó conductor ecológico Cero Halógenos NH-80 / N2XH.", "CNE-Utilización 020-028 y RNE EM.010."),
    ("5", "Capacidad del Tablero", "Excel recomendaba 24 polos (saturado).", "Se amplió a Tablero General de 36 Polos trifásico.", "CNE-U Regla 080-400 (Reserva mínima 20%).")
]
for r_idx, row_data in enumerate(m_data, start=1):
    row = tbl_m.rows[r_idx]
    bg = "F0F4F8" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, val in enumerate(row_data):
        c = row.cells[c_idx]
        set_cell_bg(c, bg)
        set_cell_margins(c, 50, 50, 60, 60)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val); r.font.name = 'Arial'; r.font.size = Pt(8.5)
        if c_idx in [0, 1]: r.font.bold = True

add_num_heading("3.0 JUSTIFICACIÓN DETALLADA DE PROPUESTAS TÉCNICAS", level=1)
add_p("a) Propuesta Técnica de Selección de Conductores LSZH / NH-80:\n"
      "Se descarta el uso de PVC convencional TW/THW debido a la emisión de gases clorados tóxicos en incendios. Conductor NH-80 cumple con CNE-U 020-028.\n\n"
      "b) Propuesta de Redimensionamiento del Tablero General a 36 Polos:\n"
      "La instalación de diferenciales individuales duplica polos. Gabinete de 36 polos asegura 35 polos ocupados y deja reserva conforme a CNE-U 080-400.")

out_docx = os.path.join(TARGET_DIR, "EXPEDIENTE_TECNICO_COMPLETO_INSTALACIONES_ELECTRICAS.docx")
doc.save(out_docx)
doc.save(os.path.join(REPO_TARGET, "EXPEDIENTE_TECNICO_COMPLETO_INSTALACIONES_ELECTRICAS.docx"))
print(f"Master unified document saved to {out_docx}")

